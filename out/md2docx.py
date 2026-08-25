# -*- coding: utf-8 -*-
"""계획 문서용 간이 md → docx 변환기 (제목/문단/표/코드블록만 지원)"""
import re, sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def clean(t):
    t = re.sub(r'<sub>(.*?)</sub>', r'\1', t)
    t = re.sub(r'&nbsp;', ' ', t)
    t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)
    t = re.sub(r'~~(.+?)~~', r'\1', t)
    t = re.sub(r'`(.+?)`', r'\1', t)
    t = re.sub(r'\[(.+?)\]\((.+?)\)', r'\1', t)
    return t.strip()

def convert(src, dst, title):
    lines = open(src, encoding='utf-8').read().split('\n')
    d = Document()
    st = d.styles['Normal']; st.font.name = '맑은 고딕'; st.font.size = Pt(9)
    h = d.add_paragraph(); r = h.add_run(title)
    r.font.size = Pt(16); r.font.bold = True; r.font.color.rgb = RGBColor(0,0x20,0x60)
    i = 0
    while i < len(lines):
        ln = lines[i]
        if ln.strip().startswith('|') and i+1 < len(lines) and re.match(r'^\|[\s:|-]+\|$', lines[i+1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [clean(c) for c in lines[i].strip().strip('|').split('|')]
                if not re.match(r'^[\s:|-]+$', '|'.join(cells)): rows.append(cells)
                i += 1
            if rows:
                n = max(len(r_) for r_ in rows)
                t = d.add_table(rows=0, cols=n); t.style = 'Light Grid Accent 1'
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                for ri, row in enumerate(rows):
                    cs = t.add_row().cells
                    for ci in range(n):
                        cs[ci].text = row[ci] if ci < len(row) else ''
                        for p in cs[ci].paragraphs:
                            for rr in p.runs:
                                rr.font.size = Pt(8)
                                if ri == 0: rr.font.bold = True
                d.add_paragraph()
            continue
        if ln.startswith('#'):
            lv = len(ln) - len(ln.lstrip('#'))
            txt = clean(ln.lstrip('#'))
            if txt:
                p = d.add_paragraph(); r = p.add_run(txt)
                r.font.bold = True
                r.font.size = Pt({1:14,2:12,3:11}.get(lv,10))
                if lv <= 2: r.font.color.rgb = RGBColor(0,0x20,0x60)
        elif ln.strip() in ('---',''):
            pass
        elif ln.strip().startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                p = d.add_paragraph(); r = p.add_run('    ' + lines[i])
                r.font.name = 'Consolas'; r.font.size = Pt(8)
                i += 1
        else:
            txt = clean(ln)
            if txt: d.add_paragraph(txt)
        i += 1
    d.save(dst); return dst

for src, dst, title in [
    ("00_skeleton_research.md", "00_뼈대_및_리서치정리.docx", "엔젤로보틱스 CB/CPS IM — 뼈대 및 리서치 정리"),
    ("01_content_plan.md",      "01_소제목별_집필계획.docx",   "엔젤로보틱스 CB/CPS IM — 소제목별 집필 계획"),
]:
    print("saved:", convert(src, dst, title))
