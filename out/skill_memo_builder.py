"""
투자심사보고서(대출·지급보증 심사) 빌더 헬퍼.

assets/template.docx 의 (투자계획서)* 스타일 체계를 그대로 사용해 문서를 생성한다.
python-docx 필요:  pip install python-docx --break-system-packages

사용 예:
    from memo_builder import Memo
    m = Memo("㈜두산 해외자회사 HyAxiom, Inc. 대출", "2026. 03.")
    m.h1("Executive Summary")
    m.body("당사는 ... 검토함.")
    m.h2("주요 금융 조건")
    m.cond_table([("차주", "HyAxiom, Inc."), ("대주", "미래에셋증권㈜")])
    m.save("out.docx")
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

TEMPLATE = Path(__file__).resolve().parent.parent / "assets" / "template.docx"

FONT_LIGHT = "KoPub돋움체_Pro Light"    # 본문 / 데이터표
FONT_MEDIUM = "KoPub돋움체_Pro Medium"  # 표 머리행 / 제목3 / 소제목
FONT_BOLD = "KoPub돋움체_Pro Bold"      # 표제 / 목차 제목 / 제목1·2
HEADER_FILL = "D9D9D9"
COVER_NAVY = "002060"   # 표제 딜명 색
ACCENT_ORANGE = "F58220"  # 표제 구분선 색
TABLE_WIDTH_DXA = 9747  # 본문 폭(11907 - 좌우여백 1080×2)에 정확히 맞춘 표 너비


def _set(el, tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn("w:" + k), str(v))
    el.append(e)
    return e


def _shade(cell, fill=HEADER_FILL):
    _set(cell._tc.get_or_add_tcPr(), "w:shd", val="clear", color="auto", fill=fill)


def _borders(cell, **sides):
    """sides: top/left/bottom/right = 'single' | 'dotted' | 'nil' | 'none'"""
    tcPr = cell._tc.get_or_add_tcPr()
    old = tcPr.find(qn("w:tcBorders"))
    if old is not None:
        tcPr.remove(old)
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        if side not in sides:
            continue
        val = sides[side]
        e = OxmlElement("w:" + side)
        e.set(qn("w:val"), val)
        if val not in ("nil", "none"):
            e.set(qn("w:sz"), "4")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "auto")
        b.append(e)
    tcPr.append(b)


def _row_height(row, twips=397):
    trPr = row._tr.get_or_add_trPr()
    _set(trPr, "w:trHeight", val=twips)


def _fill_cell(cell, text, font=FONT_LIGHT, bold=False, align=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.style = cell.part.document.styles["Normal"]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    rf.set(qn("w:eastAsia"), font)
    rf.set(qn("w:ascii"), font)
    rf.set(qn("w:hint"), "eastAsia")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _run_font(run, font, size=None, color=None, bold=True):
    run.font.name = font
    if size:
        run.font.size = Pt(size)
    rPr = run._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.insert(0, rf)
    for a in ("ascii", "eastAsia"):
        rf.set(qn("w:" + a), font)
    rf.set(qn("w:hint"), "eastAsia")
    if color:
        _set(rPr, "w:color", val=color)
    return run


def _is_number(v):
    s = str(v).strip().replace(",", "").replace("%", "").replace("+", "").replace("-", "")
    s = s.replace(".", "", 1).replace("배", "").replace("p", "")
    return s.isdigit() and s != ""


def _widths(table, widths):
    if not widths:
        return
    total = sum(widths)
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = Pt(0)  # placeholder; real width set below
            _set(cell._tc.get_or_add_tcPr(), "w:tcW",
                 w=int(TABLE_WIDTH_DXA * w / total), type="dxa")


class Memo:
    def __init__(self, deal_name="", date_line="", template=TEMPLATE, cover=True, toc=True):
        """deal_name: 딜명(표제 및 머리글에 동일하게 들어감)
        date_line : '2026. 03.' 형식
        cover/toc : 표제지·목차 자동 생성 여부"""
        self.doc = Document(str(template))
        self.deal_name = deal_name
        if deal_name:
            for p in self.doc.sections[0].header.paragraphs:
                if p.runs:
                    p.runs[0].text = deal_name
                    for r in p.runs[1:]:
                        r.text = ""
        if cover:
            self.cover(deal_name, date_line)
        if toc:
            self.toc()

    # ---------- 표제지 / 목차 ----------
    def cover(self, deal_name, date_line=""):
        """표제지. 상단 여백 → 딜명(26pt 네이비 볼드) → 주황 구분선 → 작성월(14pt 볼드) → 페이지 나눔."""
        for _ in range(6):
            self.doc.add_paragraph("", style="(투자계획서)본문")
        p = self.doc.add_paragraph(style="(투자계획서)본문")
        r = p.add_run(deal_name)
        _run_font(r, FONT_BOLD, size=26, color=COVER_NAVY)
        for _ in range(3):
            self.doc.add_paragraph("", style="(투자계획서)본문")

        # 표제 아래 주황색 짧은 구분선 (원본의 F58220 직선 도형을 표로 재현)
        rule = self.doc.add_table(rows=1, cols=1)
        _set(rule._tbl.tblPr, "w:tblW", w=1080, type="dxa")
        cell = rule.rows[0].cells[0]
        _borders(cell, top="nil", left="nil", right="nil")
        tcPr = cell._tc.get_or_add_tcPr()
        b = tcPr.find(qn("w:tcBorders"))
        bo = OxmlElement("w:bottom")
        bo.set(qn("w:val"), "single")
        bo.set(qn("w:sz"), "12")
        bo.set(qn("w:space"), "0")
        bo.set(qn("w:color"), ACCENT_ORANGE)
        b.append(bo)
        cell.paragraphs[0].style = self.doc.styles["(투자계획서)본문"]
        self.doc.add_paragraph("", style="(투자계획서)본문")

        p = self.doc.add_paragraph(style="(투자계획서)본문")
        r = p.add_run(date_line)
        _run_font(r, FONT_BOLD, size=14)
        self.page_break()

    def toc(self, levels=2, title="목차"):
        """목차 페이지. TOC 필드를 삽입하므로 Word에서 열면 자동 채워진다
        (템플릿에 updateFields=true 설정됨. 수동 갱신은 필드 선택 후 F9).

        levels=2  : 제목1·제목2까지 (원본과 동일)
        levels=3  : 제목3까지. (투자계획서)제목3 스타일에 개요 수준을 부여해야 하므로
                    호출 시 스타일을 자동 패치한다.
        """
        if levels >= 3:
            self._enable_h3_outline()
        p = self.doc.add_paragraph(style="TOC Heading")
        r = p.add_run(title)
        _run_font(r, FONT_BOLD, size=20)

        p = self.doc.add_paragraph(style="toc 1")
        r = p.add_run()
        _run_font(r, FONT_MEDIUM, size=12)
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        r._r.append(fld)
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' TOC \\o "1-{levels}" \\h \\z \\u '
        r._r.append(instr)
        sep = OxmlElement("w:fldChar")
        sep.set(qn("w:fldCharType"), "separate")
        r._r.append(sep)
        hint = OxmlElement("w:t")
        hint.text = "(Word에서 문서를 열면 목차가 자동 생성됩니다. 수동 갱신: 목차 클릭 → F9)"
        r._r.append(hint)
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        r._r.append(end)
        self.page_break()

    def _enable_h3_outline(self):
        """(투자계획서)제목3은 (투자계획서)본문 기반이라 개요 수준이 없어 기본적으로 목차에 잡히지 않는다.
        목차에 3단계를 넣으려면 개요 수준 2를 부여한다."""
        st = self.doc.styles["(투자계획서)제목3"]
        pPr = st.element.get_or_add_pPr()
        if pPr.find(qn("w:outlineLvl")) is None:
            _set(pPr, "w:outlineLvl", val="2")

    # ---------- 문단 ----------
    def _p(self, text, style):
        return self.doc.add_paragraph(str(text), style=style)

    def h1(self, text):
        """대단원. 예: Executive Summary / 리스크 분석 / 지급보증인(OOO) 분석"""
        return self._p(text, "(투자계획서)제목1")

    def appendix(self, text):
        """별첨 대단원. 제목1 서식이되 자동번호는 붙이지 않는다. 예: [별첨1] 차주 재무제표"""
        p = self._p(text, "(투자계획서)제목1")
        pPr = p._p.get_or_add_pPr()
        old = pPr.find(qn("w:numPr"))
        if old is not None:
            pPr.remove(old)
        npr = OxmlElement("w:numPr")
        for tag, val in (("w:ilvl", "0"), ("w:numId", "0")):
            e = OxmlElement(tag)
            e.set(qn("w:val"), val)
            npr.append(e)
        pPr.insert(0, npr)
        return p

    def company_profile(self, **fields):
        """기업개요표. 항목 구성과 순서를 고정해 문서 전체에서 동일하게 유지한다.
        미기재 항목은 자동으로 '-' 처리된다.

        m.company_profile(회사명="㈜두산", 본사주소="...", 주요사업="...", 대표이사="...",
                          설립일="...", 상장여부="유가증권시장 상장 / 1973년 06월 29일",
                          신용등급="BBB+ / Stable", 주요주주="박정원 외 특수관계인 (40.11%)")
        """
        order = [("회사명", "회사명"), ("본사 주소", "본사주소"), ("주요 사업", "주요사업"),
                 ("대표이사", "대표이사"), ("설립일", "설립일"),
                 ("상장 여부 / 상장일", "상장여부"), ("신용등급", "신용등급"),
                 ("주요 주주", "주요주주")]
        rows = [(label, fields.get(key, "-")) for label, key in order]
        return self.cond_table(rows, header=None, widths=(19, 81))

    def h2(self, text):
        """중단원. 예: 회사 개요 / 산업 분석 / 사업 분석 / 재무 분석"""
        return self._p(text, "(투자계획서)제목2")

    def h3(self, text):
        """소단원. 예: 수익성분석 / 유동성 리스크"""
        return self._p(text, "(투자계획서)제목3")

    def sub(self, text):
        """✓ 불릿 소제목. 예: [실질 위협 분석] / [리스크 완화 기제]"""
        return self._p(text, "(투자계획서)소제목")

    def body(self, text):
        """본문 서술. '(핵심 메시지) 근거 문장...' 형태 권장, 어미는 '~함/~임'."""
        return self._p(text, "(투자계획서)본문")

    def table_title(self, text):
        """표/그림 위 캡션. 대괄호로 감싼다. 예: [별도기준 손익계산서 요약]"""
        return self._p(text if text.startswith("[") else f"[{text}]", "(투자계획서)표 제목")

    def unit(self, text="(단위: 억원)"):
        """표 우측 상단 단위 표기."""
        return self._p(text, "(투자계획서)단위")

    def source(self, text):
        """표 아래 주석/출처. 예: '출처: ㈜두산 IR자료' / '주1) ...'"""
        return self._p(text, "(투자계획서)출처")

    def page_break(self):
        from docx.enum.text import WD_BREAK
        self.doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ---------- 표 ----------
    def cond_table(self, rows, header=("구분", "내용"), widths=(19, 81)):
        """조건표 / 개요표(2열). 좌우 외곽선 없음, 내부 점선, 머리행 음영."""
        ncol = len(header) if header else len(rows[0])
        t = self.doc.add_table(rows=0, cols=ncol)
        t.style = self.doc.styles["Table Grid"]
        _set(t._tbl.tblPr, "w:tblW", w=TABLE_WIDTH_DXA, type="dxa")
        tb = OxmlElement("w:tblBorders")
        for side, val in (("left", "none"), ("right", "none"),
                          ("insideH", "dotted"), ("insideV", "dotted")):
            e = OxmlElement("w:" + side)
            e.set(qn("w:val"), val)
            if val != "none":
                e.set(qn("w:sz"), "4")
                e.set(qn("w:space"), "0")
                e.set(qn("w:color"), "auto")
            tb.append(e)
        t._tbl.tblPr.append(tb)

        if header:
            r = t.add_row()
            _row_height(r, 397)
            for i, h in enumerate(header):
                c = r.cells[i]
                _fill_cell(c, h, font=FONT_MEDIUM, align=WD_ALIGN_PARAGRAPH.CENTER)
                _shade(c)
                _borders(c, top="single", bottom="single",
                         left="nil" if i == 0 else "dotted",
                         right="nil" if i == ncol - 1 else "dotted")
        first = header is None
        for row in rows:
            r = t.add_row()
            _row_height(r, 397)
            for i, v in enumerate(row):
                c = r.cells[i]
                _fill_cell(c, v, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else None)
                _borders(c, top="single" if first else "dotted", bottom="dotted",
                         left="nil" if i == 0 else "dotted",
                         right="nil" if i == ncol - 1 else "dotted")
            first = False
        _widths(t, widths)
        return t

    def data_table(self, header, rows, widths=None, first_col_width=26):
        """재무 데이터표. 위/아래 실선만, 머리행 음영, 항목열 우측 실선, 숫자 우측정렬."""
        ncol = len(header)
        t = self.doc.add_table(rows=0, cols=ncol)
        _set(t._tbl.tblPr, "w:tblW", w=TABLE_WIDTH_DXA, type="dxa")
        tb = OxmlElement("w:tblBorders")
        for side in ("top", "bottom"):
            e = OxmlElement("w:" + side)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "4")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "auto")
            tb.append(e)
        t._tbl.tblPr.append(tb)
        mar = OxmlElement("w:tblCellMar")
        for side in ("left", "right"):
            e = OxmlElement("w:" + side)
            e.set(qn("w:w"), "99")
            e.set(qn("w:type"), "dxa")
            mar.append(e)
        t._tbl.tblPr.append(mar)

        r = t.add_row()
        _row_height(r, 397)
        for i, h in enumerate(header):
            c = r.cells[i]
            _fill_cell(c, h, align=WD_ALIGN_PARAGRAPH.CENTER)
            _shade(c)
            _borders(c, top="single", bottom="single",
                     right="single" if i == 0 else ("nil" if i == ncol - 1 else "none"))
        for row in rows:
            r = t.add_row()
            _row_height(r, 397)
            for i, v in enumerate(row):
                c = r.cells[i]
                align = (WD_ALIGN_PARAGRAPH.LEFT if i == 0
                         else (WD_ALIGN_PARAGRAPH.RIGHT if _is_number(v)
                               else WD_ALIGN_PARAGRAPH.CENTER))
                _fill_cell(c, v, align=align)
                _borders(c, top="single", bottom="single",
                         right="single" if i == 0 else ("nil" if i == ncol - 1 else "none"))
        if widths is None:
            rest = (100 - first_col_width) / max(ncol - 1, 1)
            widths = [first_col_width] + [rest] * (ncol - 1)
        _widths(t, widths)
        return t

    def verdict_table(self, rows):
        """종합평가표. rows = [(구분, 주요내용, 시사점), ...]
        구분은 수익성 / 유동성 / 안정성 / 성장성 4행 고정."""
        return self.data_table(("구분", "주요내용", "시사점"), rows,
                               widths=(12, 44, 44))

    def save(self, path):
        self.doc.save(str(path))
        return path
